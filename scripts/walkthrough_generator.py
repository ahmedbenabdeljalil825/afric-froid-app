from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_walkthrough():
    doc = Document()

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    # --- Title Page ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Afric Froid IIoT Platform")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 32, 96) # Brand Dark Blue

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Industrial Digitalization & Real-time Monitoring")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 159, 227) # Brand Light Blue

    doc.add_page_break()

    # --- Project Objective ---
    doc.add_heading('1. Project Objective', level=1)
    doc.add_paragraph(
        "The Afric Froid IIoT Platform is a specialized industrial solution designed to bridge the gap between "
        "operational equipment and digital management. The primary objective is to provide real-time visibility, "
        "analytical insights, and remote control capabilities for industrial refrigeration and power units, "
        "ensuring operational efficiency and proactive maintenance."
    )

    # --- Key Features ---
    doc.add_heading('2. Key Features', level=1)
    features = [
        ("Real-time Telemetry Monitoring", "Live streaming of critical parameters including Temperature (°C), Suction Pressure (PSI), and Power Usage (kW)."),
        ("Interactive Equipment Control", "Ability to manage setpoints and power states remotely via secure MQTT communication."),
        ("Dynamic Widget Designer", "Advanced administrative tools to customize individual client dashboards with a variety of graphical widgets (Charts, Gauges, Status Indicators)."),
        ("Multi-tenant Architecture", "Secure user isolation allowing multiple companies to manage their own specific fleet of units."),
        ("Bilingual Interface", "Full support for English and French to cater to local and international users.")
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # --- Operational Walkthrough ---
    doc.add_heading('3. Operational Walkthrough', level=1)
    
    doc.add_heading('3.1 User Experience', level=2)
    doc.add_paragraph(
        "Upon login, clients are greeted with a customized high-fidelity dashboard. The interface provides "
        "instant visual feedback on the 'PLC Online' status and live telemetry trends. The 'Controls' section "
        "enables authorized users to adjust operational parameters with immediate effect on the field equipment."
    )

    doc.add_heading('3.2 Administrative Control', level=2)
    doc.add_paragraph(
        "Administrators have access to a powerful 'Widget Designer'. This interface allows them to drag, drop, "
        "and configure widgets for any specific user. They can define MQTT topics, JSON extraction paths, "
        "and visualization types (e.g., Line Charts vs. Gauges), enabling a bespoke monitoring experience for "
        "different industrial use cases."
    )

    # --- Technical Architecture ---
    doc.add_heading('4. Technical Architecture', level=1)
    doc.add_paragraph(
        "The platform is built on a modern, scalable stack designed for high availability and low latency:"
    )
    tech_stack = [
        ("Core Stack", "Built with React and TypeScript for a robust, type-safe frontend experience."),
        ("Cloud Infrastructure", "Powered by Supabase for secure authentication and real-time database management."),
        ("Connectivity Layer", "Utilizes the MQTT protocol, the industry standard for IoT, to ensure lightweight and reliable messaging between the web platform and industrial PLCs."),
        ("UI Framework", "Custom-styled using Vanilla CSS and Tailwind for a premium, responsive industrial aesthetic.")
    ]
    for title, desc in tech_stack:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # --- Conclusion ---
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "The Afric Froid IIoT Platform represents a significant step forward in industrial monitoring. By "
        "combining real-time connectivity with a highly flexible user interface, it provides the tools "
        "necessary for modern industrial teams to optimize their operations and maintain peak performance."
    )

    # Save the document
    output_path = "App_Walkthrough.docx"
    doc.save(output_path)
    print(f"Document saved successfully to {os.path.abspath(output_path)}")

if __name__ == "__main__":
    create_walkthrough()
